#!/usr/bin/env python3
"""
One-time helper: pull every plausible bullet / summary line from each known
resume DOCX so we can hand-curate `bullets.yaml`.

Output is written to stdout as a markdown-ish dump grouped by file. The
downstream step is human: WS Gong verifies every line before we put it in
`bullets.yaml` (see CLAUDE.md §2 Phase 2).
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parent.parent


def collect_sources() -> list[Path]:
    """Discover resume DOCX inputs at runtime.

    Hardcoding paths (NVIDIA/, Vercel/, etc.) went stale after CLAUDE.md §10
    Phase 5 moved legacy company folders under applications/. Glob instead.
    """
    out: list[Path] = [REPO / "resume-template.docx"]
    # Generalized resume (date-prefixed filenames at the root).
    out.extend(sorted(REPO.glob("*-wsgong-resume-generalized.docx")))
    apps = REPO / "applications"
    if apps.exists():
        # Tailored resumes per application — skip the _template skeleton.
        for pattern in ("resume.docx",
                        "WSGong_Resume_*.docx",
                        "billy-gong-resume-*.docx"):
            for path in sorted(apps.rglob(pattern)):
                if "_template" in path.parts:
                    continue
                out.append(path)
    return out


def iter_text(doc: Document):
    """Walk every paragraph in the main doc body, then every table cell.

    The Swiss template keeps body content in paragraphs plus a 5-row table,
    so we have to look in both places.
    """
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            yield ("para", text)

    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        yield (f"table[{row_idx},{cell_idx}]", text)


def dump(path: Path) -> None:
    print(f"\n\n===== {path.relative_to(REPO)} =====")
    if not path.exists():
        print("(missing)")
        return
    doc = Document(str(path))
    seen = set()
    for location, text in iter_text(doc):
        # Collapse exact duplicates per-file (the template has repeated shells).
        key = (location, text)
        if key in seen:
            continue
        seen.add(key)
        print(f"- [{location}] {text}")


def main() -> int:
    sources = collect_sources()
    if not sources:
        print("extract_bullets.py: no DOCX sources discovered.")
        return 1
    for path in sources:
        dump(path)
    return 0


if __name__ == "__main__":
    sys.exit(main())
