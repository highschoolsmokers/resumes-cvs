#!/usr/bin/env python3
"""Detect bullets that may have been hand-edited on a tailored resume.

CURRENT STATUS (detection only):
    This script identifies bullet IDs whose `bullets.yaml` text does NOT
    appear verbatim in the tailored `resume.docx`. It surfaces those as
    candidates the user should review. It does **not** rewrite
    `bullets.yaml` — preserving the file's schema-documenting comments
    through programmatic YAML round-trips is non-trivial, and we'd rather
    print proposals than silently corrupt the closed universe.

    The interactive `[u]pdate` / `[n]ew-id` prompt collects choices but
    `save_bullets()` is currently a stub that prints the proposals and
    tells the user to apply them by hand. Implementing the in-place
    rewrite (likely via ruamel.yaml or string surgery on the YAML body)
    is a follow-up.

Problem the full feature solves: the user tailors a resume, then hand-edits
the generated `resume.docx` — sharpening a bullet, fixing a typo, tightening
phrasing. Next time the tailor runs, those edits would be lost because the
source of truth is `bullets.yaml`, not the tailored docx.

Mapping rendered text back to bullet IDs uses `resume.provenance.yaml` —
every claim in the sidecar names its source ID, so the mapping is explicit.

Usage:
    python scripts/backprop_edits.py applications/Anthropic/fde-2026-04-20/
    python scripts/backprop_edits.py --dry-run applications/Anthropic/fde-2026-04-20/

Both modes today produce the same output: a list of suspect bullet IDs and
the original text. `--dry-run` skips the interactive prompt loop.
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    import yaml  # type: ignore
    from docx import Document  # type: ignore
except ImportError as e:  # pragma: no cover
    sys.stderr.write(f"backprop_edits.py needs PyYAML and python-docx: {e}\n")
    sys.exit(1)


REPO = Path(__file__).resolve().parent.parent
BULLETS_YAML = REPO / "bullets.yaml"


def load_bullets() -> dict:
    with BULLETS_YAML.open() as f:
        return yaml.safe_load(f)


def save_bullets(data: dict) -> None:  # noqa: ARG001
    """Placeholder — see module docstring.

    Implementing in-place rewrite of bullets.yaml requires preserving the
    file's schema-documenting comments, which `yaml.safe_dump` drops. A real
    implementation would either:
      (a) switch to ruamel.yaml round-trip mode, or
      (b) do targeted string surgery on the `bullets:` block.

    Until that lands, the user applies the proposed edits by hand.
    """
    sys.stderr.write(
        "backprop_edits.py: in-place bullets.yaml rewrite is not implemented.\n"
        "  Apply the proposed edits above by hand to preserve comments.\n"
    )


def extract_experience_bullets(resume_docx: Path) -> dict[str, str]:
    """Map heading → body text from the Experience section of a resume DOCX.

    We rely on the Swiss layout: Experience entries are rows in a table whose
    left cell carries dates and whose right cell's first paragraph is
    "{employer} — {title}" and second paragraph is the bullet body.
    """
    doc = Document(str(resume_docx))
    out: dict[str, str] = {}
    in_experience = False
    for table in doc.tables:
        for row in table.rows:
            if not row.cells:
                continue
            right = row.cells[-1]
            paragraphs = [p.text.strip() for p in right.paragraphs if p.text.strip()]
            if not paragraphs:
                continue
            if paragraphs[0] == "Experience":
                in_experience = True
                continue
            if paragraphs[0] in {"Education", "Publications & Activity",
                                 "Community", "Projects", "Skills"}:
                in_experience = False
                continue
            if in_experience and len(paragraphs) >= 2:
                head, body = paragraphs[0], " ".join(paragraphs[1:])
                out[head] = body
    return out


def load_provenance(app_folder: Path) -> dict:
    sidecar = app_folder / "resume.provenance.yaml"
    if not sidecar.exists():
        sys.stderr.write(
            f"backprop_edits.py: no sidecar at {sidecar}. "
            "Can't map rendered bullets to source IDs.\n"
        )
        sys.exit(1)
    with sidecar.open() as f:
        return yaml.safe_load(f) or {}


def prompt(question: str, choices: str) -> str:
    while True:
        sys.stderr.write(f"{question} [{choices}] ")
        sys.stderr.flush()
        resp = sys.stdin.readline().strip().lower()
        if resp and resp in choices:
            return resp


def normalize(text: str) -> str:
    """Collapse whitespace for fuzzy comparison."""
    return re.sub(r"\s+", " ", text or "").strip()


def propose_changes(app_folder: Path) -> list[dict]:
    """Return list of {bullet_id, original, edited, action_default}."""
    sidecar = load_provenance(app_folder)
    bullets_data = load_bullets()
    bullets_by_id = {b["id"]: b for b in bullets_data.get("bullets", [])}

    resume = app_folder / "resume.docx"
    if not resume.exists():
        sys.stderr.write(f"backprop_edits.py: {resume} not found.\n")
        sys.exit(1)
    experience = extract_experience_bullets(resume)

    # Build a flat list of rendered bullets in the order they appeared, with
    # the IDs the sidecar says they map to. The sidecar uses `source:
    # bullet:<id>` entries — one per claim.
    by_id: dict[str, str] = {}
    for claim in sidecar.get("claims", []) or []:
        source = (claim or {}).get("source") or ""
        if not source.startswith("bullet:"):
            continue
        bid = source.split(":", 1)[1]
        # Best-effort match — the bullet's rendered `text` should appear in
        # the combined Experience body strings.
        by_id.setdefault(bid, "")

    # For each bullet_id in the sidecar, compare original vs. rendered.
    proposals: list[dict] = []
    for bid, _ in by_id.items():
        original = bullets_by_id.get(bid, {}).get("text", "")
        if not original:
            continue
        rendered = None
        for body in experience.values():
            if normalize(original) in normalize(body):
                rendered = normalize(original)
                break
        # If we can't find the original verbatim, the user most likely edited
        # it. Try to locate the role in experience and extract the diverged
        # sentence that most closely matches.
        if rendered is None:
            role_id = bullets_by_id[bid].get("role_id")
            # Fall back: flag as possibly-edited and show full experience body.
            proposals.append({
                "bullet_id": bid,
                "original": original,
                "edited": "[MANUAL EDIT SUSPECTED — review "
                          f"resume.docx role {role_id!r} by hand]",
            })
    return proposals


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__,
                                     formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("app_folder", type=Path,
                        help="application folder to back-propagate from")
    parser.add_argument("--dry-run", action="store_true",
                        help="only print proposals; do not prompt")
    args = parser.parse_args()

    if not args.app_folder.is_dir():
        sys.stderr.write(f"backprop_edits.py: not a directory: {args.app_folder}\n")
        return 1

    proposals = propose_changes(args.app_folder)

    if not proposals:
        print("backprop_edits.py: no divergence detected. "
              "Every source bullet appears verbatim in the resume.")
        return 0

    print(f"backprop_edits.py: {len(proposals)} bullet(s) may have been edited:\n")
    for p in proposals:
        print(f"- {p['bullet_id']}")
        print(f"    original: {p['original']}")
        print(f"    edited:   {p['edited']}")
        print()

    if args.dry_run:
        return 0

    # Until the in-place save is implemented, the interactive loop only
    # confirms intent and emits the proposed edits — the user applies them.
    approved: list[dict] = []
    for p in proposals:
        print(f"\nBullet: {p['bullet_id']}")
        print(f"  original: {p['original']}")
        print(f"  edited:   {p['edited']}")
        choice = prompt("Mark for update [u], new-id [n], skip [s], quit [q]?", "unsq")
        if choice == "q":
            break
        if choice == "s":
            continue
        if choice == "u":
            approved.append({**p, "action": "update"})
        elif choice == "n":
            new_id = input("  new bullet id (kebab-case): ").strip()
            approved.append({**p, "action": "new", "new_id": new_id})

    if not approved:
        print("No changes marked. Done.")
        return 0

    print("\nMarked for edit:")
    for a in approved:
        print(f"  - {a}")

    save_bullets({})  # stub — prints "apply by hand"; see module docstring.
    return 0


if __name__ == "__main__":
    sys.exit(main())
