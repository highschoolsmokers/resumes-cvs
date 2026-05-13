#!/usr/bin/env python3
"""Render `cover-letter.md` → `cover-letter.docx` with the W.S. Gong letterhead.

The docx output then goes through `scripts/docx_to_pdf.py` to produce the
PDF the user attaches to applications. Typography mirrors the resume's
Swiss style: Raleway display, Lato body, #D44500 accent.

Input markdown format (produced by agents/cover-letter-writer.md):

    Dear <salutation>,

    <Hook paragraph>

    <Bridge paragraph>

    <Close paragraph>

    W.S. Gong

Blank lines separate paragraphs. The first non-empty line is treated as
the salutation; the last non-empty line is treated as the signature.
Everything in between is body.

Guardrails enforced at render time (spec §5.4, §9.3):

    - Word count of the body (excluding salutation and signature) must
      be ≤ `voice.yaml → length.hard_max` (default 500). If it isn't,
      the script refuses to render and tells the user which bound was
      crossed.
    - Any hit against `voice.yaml → forbidden_phrases` prints a WARNING
      but does not abort. The writer agent's self-audit is the first
      line of defense; this is the second.
    - The presence of any `[NEEDS SOURCE: ...]` marker aborts the
      render — unsourced claims must be resolved before the document
      is handed to the user.

Usage:

    python3 build_cover_letter.py \
        --input applications/<...>/cover-letter.md \
        --out   applications/<...>/cover-letter.docx
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    import yaml  # type: ignore
except ImportError:  # pragma: no cover
    sys.stderr.write("build_cover_letter.py needs PyYAML. pip install -r requirements.txt\n")
    sys.exit(1)

try:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.shared import Pt, RGBColor, Inches  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
except ImportError:  # pragma: no cover
    sys.stderr.write("build_cover_letter.py needs python-docx. pip install -r requirements.txt\n")
    sys.exit(1)


REPO = Path(__file__).resolve().parent
VOICE_YAML = REPO / "config" / "voice.yaml"

# Style constants — mirrors docs/resume-style-spec.md.
ACCENT = RGBColor(0xD4, 0x45, 0x00)      # #D44500
INK    = RGBColor(0x1A, 0x1A, 0x1A)      # near-black body
MUTED  = RGBColor(0x66, 0x66, 0x66)      # subheadings, contact line

DISPLAY_FONT = "Raleway"
BODY_FONT    = "Lato"


# ─── Voice config ──────────────────────────────────────────────────────

def load_voice(path: Path = VOICE_YAML) -> dict:
    if not path.exists():
        sys.stderr.write(f"build_cover_letter.py: missing {path}\n")
        sys.exit(2)
    with path.open() as f:
        return yaml.safe_load(f) or {}


# ─── Parse markdown ────────────────────────────────────────────────────

def parse_letter_md(path: Path) -> tuple[str, list[str], str]:
    """Return (salutation, body_paragraphs, signature).

    Paragraphs are whatever lives between blank lines, joined on a single
    space so soft-wrapped markdown becomes flowing prose.
    """
    raw = path.read_text(encoding="utf-8")

    # Split into blocks separated by blank lines.
    blocks: list[str] = []
    buf: list[str] = []
    for line in raw.splitlines():
        if line.strip() == "":
            if buf:
                blocks.append(" ".join(b.strip() for b in buf).strip())
                buf = []
        else:
            buf.append(line)
    if buf:
        blocks.append(" ".join(b.strip() for b in buf).strip())

    blocks = [b for b in blocks if b]
    if len(blocks) < 3:
        sys.stderr.write(
            f"build_cover_letter.py: expected at least "
            f"salutation + 1 paragraph + signature; got {len(blocks)} blocks\n"
        )
        sys.exit(2)

    salutation = blocks[0]
    signature = blocks[-1]
    body = blocks[1:-1]
    return salutation, body, signature


# ─── Guardrails ────────────────────────────────────────────────────────

WORD_RE = re.compile(r"\w[\w'-]*", re.UNICODE)


def word_count(text: str) -> int:
    return len(WORD_RE.findall(text))


def check_needs_source(salutation: str, body: list[str], signature: str) -> None:
    joined = "\n".join([salutation, *body, signature])
    if "[NEEDS SOURCE" in joined:
        sys.stderr.write(
            "build_cover_letter.py: refusing to render — draft contains "
            "[NEEDS SOURCE: ...] placeholder(s). Resolve before building.\n"
        )
        sys.exit(2)


def check_length(body: list[str], voice: dict) -> int:
    """Return body word count. Abort if over hard_max."""
    count = sum(word_count(p) for p in body)
    length_cfg = voice.get("length") or {}
    hard_max = int(length_cfg.get("hard_max", 500))
    target_min = int(length_cfg.get("target_min", 300))
    target_max = int(length_cfg.get("target_max", 400))

    if count > hard_max:
        sys.stderr.write(
            f"build_cover_letter.py: body is {count} words; "
            f"voice.yaml hard_max is {hard_max}. Cut and rerun.\n"
        )
        sys.exit(2)
    if count < target_min:
        sys.stderr.write(
            f"build_cover_letter.py: WARNING — body is {count} words; "
            f"target_min is {target_min}. Consider adding one specific detail.\n"
        )
    elif count > target_max:
        sys.stderr.write(
            f"build_cover_letter.py: WARNING — body is {count} words; "
            f"target_max is {target_max} (hard_max {hard_max}). "
            "Consider cutting the weakest evidence point.\n"
        )
    return count


def check_forbidden(body: list[str], voice: dict) -> list[str]:
    """Return list of forbidden phrases found. Warn-only; do not abort."""
    phrases = voice.get("forbidden_phrases") or []
    hits: list[str] = []
    joined_lower = "\n".join(body).lower()
    for phrase in phrases:
        if phrase.lower() in joined_lower:
            hits.append(phrase)
    for phrase in hits:
        sys.stderr.write(
            f"build_cover_letter.py: WARNING — forbidden_phrase '{phrase}' "
            "appears in the draft. Rewrite.\n"
        )
    return hits


# ─── DOCX rendering ────────────────────────────────────────────────────

def _set_run_font(run, font_name: str, size_pt: float,
                  color: RGBColor | None = None, bold: bool = False) -> None:
    run.font.name = font_name
    # East-Asian fallback — matches the resume template's behaviour so the
    # font doesn't fall back to Calibri on Word-for-Windows.
    rFonts = run._element.rPr.rFonts if run._element.rPr is not None else None
    if rFonts is None:
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn(f"w:{attr}"), font_name)
    run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color
    if bold:
        run.bold = True


def _thin_rule(doc, color: RGBColor = ACCENT) -> None:
    """Insert a thin horizontal rule via a 1pt bottom border on an empty paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")           # 1/8 pt units; 6 = 0.75pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"),
               f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _letterhead(doc, contact: dict) -> None:
    # Name: large Raleway, accent color.
    name_para = doc.add_paragraph()
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after = Pt(0)
    name_run = name_para.add_run(contact.get("name", "W.S. Gong"))
    _set_run_font(name_run, DISPLAY_FONT, 26, ACCENT, bold=True)

    # Subhead: smaller Raleway, muted.
    sub_para = doc.add_paragraph()
    sub_para.paragraph_format.space_before = Pt(0)
    sub_para.paragraph_format.space_after = Pt(2)
    sub_run = sub_para.add_run(contact.get("subhead", "Engineer, Editor, Writer"))
    _set_run_font(sub_run, DISPLAY_FONT, 11, MUTED, bold=False)

    # Contact line: one paragraph, Lato small, muted.
    parts = contact.get("contact_line", [
        "San Francisco, CA",
        "billygong@me.com",
        "ws-gong.com",
        "github.com/highschoolsmokers",
    ])
    contact_para = doc.add_paragraph()
    contact_para.paragraph_format.space_before = Pt(0)
    contact_para.paragraph_format.space_after = Pt(0)
    sep = "  ·  "
    contact_run = contact_para.add_run(sep.join(parts))
    _set_run_font(contact_run, BODY_FONT, 9, MUTED, bold=False)

    _thin_rule(doc, ACCENT)


def _body_paragraph(doc, text: str, *, first_of_body: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6 if first_of_body else 8)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    _set_run_font(run, BODY_FONT, 10.5, INK, bold=False)


def _simple_line(doc, text: str, *, before: int = 8, after: int = 0,
                 size: float = 10.5, italic: bool = False,
                 align=None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _set_run_font(run, BODY_FONT, size, INK, bold=False)
    if italic:
        run.italic = True


def render_letter(
    md_path: Path,
    out_path: Path,
    voice: dict,
    date_override: str | None = None,
) -> None:
    salutation, body, signature = parse_letter_md(md_path)

    # Guardrails — abort BEFORE writing the docx if the draft is broken.
    check_needs_source(salutation, body, signature)
    words = check_length(body, voice)
    check_forbidden(body, voice)

    doc = Document()

    # Page margins — match resume (~0.6in all around, a hair wider at top).
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Default Normal style → Lato 10.5 so any stray paragraph is legible.
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK

    contact = voice.get("letterhead", {}) or {}
    _letterhead(doc, contact)

    # Date — right-aligned, muted Lato.
    import datetime as _dt
    date_str = date_override or _dt.date.today().strftime("%B %-d, %Y") \
        if sys.platform != "win32" else (date_override or _dt.date.today().strftime("%B %d, %Y"))
    _simple_line(doc, date_str, before=10, after=2, size=10,
                 align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Salutation.
    _simple_line(doc, salutation, before=12, after=0, size=10.5)

    # Body paragraphs.
    for i, paragraph in enumerate(body):
        _body_paragraph(doc, paragraph, first_of_body=(i == 0))

    # Signature line — one blank gap, then the name.
    _simple_line(doc, "", before=4, after=0, size=10.5)
    _simple_line(doc, signature, before=10, after=0, size=10.5)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

    sys.stderr.write(
        f"Cover letter written → {out_path}  ({words} body words, "
        f"{len(body)} paragraph(s))\n"
    )


# ─── CLI ───────────────────────────────────────────────────────────────

def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__,
                                     formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--input", type=Path, required=True,
                        help="path to cover-letter.md")
    parser.add_argument("--out", type=Path, required=True,
                        help="path to emit cover-letter.docx")
    parser.add_argument("--date", type=str, default=None,
                        help="override the date line (e.g. 'April 20, 2026')")
    args = parser.parse_args()

    if not args.input.exists():
        sys.stderr.write(f"build_cover_letter.py: {args.input} not found\n")
        return 2

    voice = load_voice()
    render_letter(args.input, args.out, voice, date_override=args.date)
    return 0


if __name__ == "__main__":
    sys.exit(main())
